#!/usr/bin/env python3
import sys
import requests
import os
import tempfile
from urllib.parse import urlparse
import json
import subprocess

# Import processing libraries
try:
    import PyPDF2
    from docx import Document
    import pdfplumber
except ImportError as e:
    print(f"Import error: {e}")
    sys.exit(1)

def download_file(url, filename):
    """Download file from URL"""
    try:
        response = requests.get(url, stream=True, timeout=30)
        response.raise_for_status()
        
        with open(filename, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        return True
    except Exception as e:
        print(f"Download error: {e}")
        return False

def extract_text_from_pdf(filepath):
    """Extract text from PDF using multiple methods"""
    text = ""
    
    # Method 1: Try pdfplumber (best for complex layouts)
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return clean_text(text)
    except Exception as e:
        print(f"pdfplumber failed: {e}")
    
    # Method 2: Try PyPDF2
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return clean_text(text)
    except Exception as e:
        print(f"PyPDF2 failed: {e}")
    
    # Method 3: Try pdftotext (command line)
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', filepath, '-'],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return clean_text(result.stdout)
    except Exception as e:
        print(f"pdftotext failed: {e}")
    
    return ""

def extract_text_from_docx(filepath):
    """Extract text from DOCX file"""
    try:
        doc = Document(filepath)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return clean_text(text)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_doc(filepath):
    """Extract text from DOC file using antiword"""
    try:
        result = subprocess.run(
            ['antiword', filepath],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0:
            return clean_text(result.stdout)
    except Exception as e:
        print(f"DOC extraction failed: {e}")
    return ""

def clean_text(text):
    """Clean and normalize extracted text"""
    import re
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\-\.,;:()\[\]]+', '', text)
    # Trim
    text = text.strip()
    
    return text

def process_file(file_url, application_id, callback_url, auth_token):
    """Main processing function"""
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_filename = tmp_file.name
        
        # Get file extension from URL
        parsed_url = urlparse(file_url)
        file_extension = os.path.splitext(parsed_url.path)[1].lower()
        
        if not file_extension:
            # Try to get from Content-Type header
            try:
                head_response = requests.head(file_url, timeout=10)
                content_type = head_response.headers.get('content-type', '').lower()
                if 'pdf' in content_type:
                    file_extension = '.pdf'
                elif 'wordprocessingml' in content_type:
                    file_extension = '.docx'
                elif 'msword' in content_type:
                    file_extension = '.doc'
            except:
                pass
        
        # Add extension to temp file
        if file_extension:
            tmp_filename += file_extension
    
    try:
        # Download file
        print(f"Downloading file from: {file_url}")
        if not download_file(file_url, tmp_filename):
            raise Exception("Failed to download file")
        
        # Determine file type and extract text
        extracted_text = ""
        
        if file_extension == '.pdf':
            print("Processing PDF file...")
            extracted_text = extract_text_from_pdf(tmp_filename)
        elif file_extension == '.docx':
            print("Processing DOCX file...")
            extracted_text = extract_text_from_docx(tmp_filename)
        elif file_extension == '.doc':
            print("Processing DOC file...")
            extracted_text = extract_text_from_doc(tmp_filename)
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
        
        if not extracted_text.strip():
            raise Exception("No text could be extracted from the file")
        
        # Prepare result
        result = {
            "success": True,
            "application_id": application_id,
            "extracted_text": extracted_text,
            "file_type": file_extension,
            "text_length": len(extracted_text),
            "processing_info": {
                "method": "github_actions",
                "timestamp": requests.get("http://worldtimeapi.org/api/ip").json()["utc_datetime"]
            }
        }
        
        print(f"Text extraction successful. Length: {len(extracted_text)} characters")
        print(f"First 200 characters: {extracted_text[:200]}...")
        
    except Exception as e:
        result = {
            "success": False,
            "application_id": application_id,
            "error": str(e),
            "processing_info": {
                "method": "github_actions",
                "timestamp": requests.get("http://worldtimeapi.org/api/ip").json()["utc_datetime"]
            }
        }
        print(f"Processing failed: {e}")
    
    finally:
        # Clean up temporary file
        if os.path.exists(tmp_filename):
            os.unlink(tmp_filename)
    
    # Send result back to Laravel app
    try:
        print(f"Sending result to callback URL: {callback_url}")
        response = requests.post(
            callback_url,
            json=result,
            headers={
                'Authorization': f'Bearer {auth_token}',
                'Content-Type': 'application/json'
            },
            timeout=30
        )
        response.raise_for_status()
        print("Callback successful")
    except Exception as e:
        print(f"Callback failed: {e}")
        # Try to log the failure (you could send to another endpoint)
        
    return result

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: process_cv.py <file_url> <application_id> <callback_url> <auth_token>")
        sys.exit(1)
    
    file_url = sys.argv[1]
    application_id = sys.argv[2]
    callback_url = sys.argv[3]
    auth_token = sys.argv[4]
    
    result = process_file(file_url, application_id, callback_url, auth_token)
    
    if result["success"]:
        print("✅ Processing completed successfully")
        sys.exit(0)
    else:
        print("❌ Processing failed")
        sys.exit(1)
